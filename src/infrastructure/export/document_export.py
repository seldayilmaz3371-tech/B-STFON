"""
Basit markdown → .docx dönüştürücü.

KAPSAM (bilinçli sınırlama): YALNIZCA "## Başlık" biçimli markdown
başlıklarını ve düz paragrafları destekler — AIInsightService.
generate_portfolio_report()'un ÜRETTİĞİ format İLE TUTARLI (o metod
"## Genel Bakış" gibi başlıklar üretmesi için AÇIKÇA talimatlandırıldı).
Listeler, tablolar, kalın/italik gibi ZENGİN markdown biçimlendirmesi
DESTEKLENMİYOR — bu, genel amaçlı bir markdown→docx dönüştürücü DEĞİL,
AI raporlarının SPESİFİK, BASİT formatı için yeterli bir araç.

NEDEN python-docx (npm docx-js DEĞİL): Bu proje boyunca ARTIFACT
oluşturmak için kullanılan docx-js (Node.js) yalnızca BENİM kendi
araç setimde çalışıyor — DAĞITILAN Streamlit uygulamasının çalışma
zamanında Node.js bulunması GARANTİ DEĞİL. python-docx, saf Python
bir kütüphane — uygulamanın KENDİ Python ortamında (requirements.txt
üzerinden) çalışır, ek bir çalışma zamanı bağımlılığı (Node.js)
GEREKTİRMEZ.
"""

from __future__ import annotations

from io import BytesIO


def markdown_report_to_docx_bytes(markdown_text: str, title: str) -> bytes:
    """
    Args:
        markdown_text: "## Başlık" biçimli basit markdown metni
            (AIInsightService.generate_portfolio_report()'un çıktısı).
        title: Belgenin ana başlığı (portföy adı gibi).

    Returns:
        .docx dosyasının ham byte içeriği — Streamlit'in
        st.download_button(data=...) parametresine DOĞRUDAN verilebilir.
    """
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=0)

    for raw_line in markdown_text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        else:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
