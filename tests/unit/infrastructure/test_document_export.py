"""markdown_report_to_docx_bytes() testleri — gerçek .docx üretip geri okuyarak doğrular."""

from __future__ import annotations

from io import BytesIO

import pytest

from src.infrastructure.export.document_export import markdown_report_to_docx_bytes


def test_produces_valid_docx_bytes():
    """Üretilen byte'lar GERÇEKTEN geçerli bir .docx olmalı — python-docx ile geri açılabilmeli."""
    from docx import Document

    result = markdown_report_to_docx_bytes("## Başlık\n\nİçerik.", "Test Raporu")
    doc = Document(BytesIO(result))  # ÇÖKMEMELİ — geçersiz .docx olsaydı burada patlardı
    assert len(doc.paragraphs) > 0


def test_title_becomes_document_title():
    from docx import Document

    result = markdown_report_to_docx_bytes("İçerik.", "Portföyüm Raporu")
    doc = Document(BytesIO(result))
    assert doc.paragraphs[0].text == "Portföyüm Raporu"
    assert doc.paragraphs[0].style.name == "Title"


def test_double_hash_becomes_heading_2():
    from docx import Document

    result = markdown_report_to_docx_bytes("## Risk Değerlendirmesi\n\nMetin.", "Rapor")
    doc = Document(BytesIO(result))
    heading_paragraphs = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert len(heading_paragraphs) == 1
    assert heading_paragraphs[0].text == "Risk Değerlendirmesi"


def test_regular_lines_become_normal_paragraphs():
    from docx import Document

    result = markdown_report_to_docx_bytes("Bu normal bir paragraf.", "Rapor")
    doc = Document(BytesIO(result))
    normal_paragraphs = [p for p in doc.paragraphs if p.text == "Bu normal bir paragraf."]
    assert len(normal_paragraphs) == 1
    assert normal_paragraphs[0].style.name == "Normal"


def test_empty_lines_are_skipped_not_added_as_blank_paragraphs():
    from docx import Document

    result = markdown_report_to_docx_bytes("Satır 1\n\n\nSatır 2", "Rapor")
    doc = Document(BytesIO(result))
    texts = [p.text for p in doc.paragraphs]
    assert "" not in texts  # boş paragraf EKLENMEMİŞ olmalı


def test_multiple_headings_and_paragraphs_preserve_order():
    """KRİTİK doğrulama: belge, GİRİŞ metnindeki SIRAYI korumalı — karışmamalı."""
    from docx import Document

    markdown = "## Bölüm 1\n\nBirinci içerik.\n\n## Bölüm 2\n\nİkinci içerik."
    result = markdown_report_to_docx_bytes(markdown, "Rapor")
    doc = Document(BytesIO(result))

    texts_in_order = [p.text for p in doc.paragraphs if p.text.strip()]
    assert texts_in_order == ["Rapor", "Bölüm 1", "Birinci içerik.", "Bölüm 2", "İkinci içerik."]
